from __future__ import annotations

 

import argparse

import gc

import hashlib

import os

import re

import time

import warnings

from bisect import bisect_right

from dataclasses import dataclass

from pathlib import Path

from collections import defaultdict

from typing import Dict, Iterable, List, Optional, Set, Tuple

 

try:

    from openpyxl import Workbook, load_workbook

    from openpyxl.drawing.image import Image as XLImage

    from openpyxl.utils.cell import coordinate_to_tuple, get_column_letter, range_boundaries

except Exception:

    Workbook = None

    XLImage = None

    load_workbook = None

    coordinate_to_tuple = None

    get_column_letter = None

    range_boundaries = None

 

try:

    from docx import Document

    from docx.oxml.ns import qn

except Exception:

    Document = None

    qn = None

 

try:

    from pypdf import PdfReader

except Exception:

    PdfReader = None

 

try:

    import fitz

except Exception:

    fitz = None

 

try:

    from win32com.client import gencache

except Exception:

    gencache = None

 

MM_PER_INCH = 25.4

EMUS_PER_INCH = 914400

EMU_PER_MM = EMUS_PER_INCH / MM_PER_INCH

POINTS_PER_INCH = 72.0

PIXELS_PER_INCH = 96.0

COMMON_MARGIN_RULE_TEXT = "余白（A4縦:上20/下20/左30/右20mm以上）"

 

@dataclass

class CheckResult:

    file_path: str

    file_type: str

    check_id: str

    check_item: str

    status: str  # PASS / FAIL / WARN / MANUAL / ERROR / N/A

    detail: str

    suggested_action: str

 

@dataclass

class VisualPage:

    file_path: str

    page_no: int

    image_path: str

    sheet_name: str = ""

 

COMMON_CHECK_ITEMS: List[Tuple[str, str]] = [

    ("C1", "プロパティ情報削除"),

    ("C2", "表紙が規定のもの"),

    ("C3", COMMON_MARGIN_RULE_TEXT),

    ("C4", "ページ番号"),

    ("C5", "PDF出力結果確認（見切れ/罫線/表サイズ/ページ番号）"),

]

 

TYPE_SPECIFIC_CHECK_ITEMS = {

    "Excel": [

        ("E1", "印刷範囲外記載チェック"),

        ("E2", "不要シートチェック"),

        ("E3", "印刷プレビュー（見切れ/罫線欠け）"),

        ("E4", "見え消し（取り消し線/訂正線）残存"),

        ("E5", "コメント（吹き出し）残存"),

        ("E6", "参照エラー残存"),

        ("E7", "ページ設定のタイトル行"),

    ],

    "Word": [

        ("W1", "見え消し（変更履歴）残存"),

        ("W2", "マーカ残存"),

        ("W3", "不要な文字色"),

        ("W4", "コメント（吹き出し）残存"),

        ("W5", "参照エラー残存"),

    ],

    "PDF": [

        ("P1", "見え消し（注釈ベース）残存"),

        ("P2", "コメント（吹き出し）残存"),

        ("P3", "参照エラー残存"),

    ],

    "LegacyOffice": [("L1", "旧形式ファイル")],

}

 

def inches_to_mm(value: Optional[float]) -> Optional[float]:

    if value is None:

        return None

    try:

        return float(value) * MM_PER_INCH

    except Exception:

        return None

 

def mm_to_points(value_mm: float) -> float:

    return float(value_mm) * POINTS_PER_INCH / MM_PER_INCH

 

def add_result(

    results: List[CheckResult],

    file_path: Path,

    file_type: str,

    check_id: str,

    check_item: str,

    status: str,

    detail: str,

    suggested_action: str,

) -> None:

    results.append(

        CheckResult(

            file_path=str(file_path),

            file_type=file_type,

            check_id=check_id,

            check_item=check_item,

            status=status,

            detail=detail,

            suggested_action=suggested_action,

        )

    )

 

def ensure_expected_checks(

    results: List[CheckResult],

    file_path: Path,

    file_type: str,

) -> None:

    target = str(file_path)

    present_ids = {r.check_id for r in results if r.file_path == target}

    expected = TYPE_SPECIFIC_CHECK_ITEMS.get(file_type, []) + COMMON_CHECK_ITEMS

    for check_id, check_item in expected:

        if check_id in present_ids:

            continue

        add_result(

            results,

            file_path,

            file_type,

            check_id,

            check_item,

            "N/A",

            "このファイルでは自動判定未対応、または判定対象外です。",

            "必要に応じて手動確認してください。",

        )

 

def slugify_for_path(path: Path) -> str:

    base = re.sub(r"[^A-Za-z0-9_.-]+", "_", path.stem)

    parent = re.sub(r"[^A-Za-z0-9_.-]+", "_", path.parent.name or "root")

    short_parent = parent[:16] or "root"

    short_base = base[:24] or "file"

    hash8 = hashlib.sha1(str(path).encode("utf-8", errors="ignore")).hexdigest()[:8]

    return f"{short_parent}_{short_base}_{hash8}"

 

def convert_office_to_pdf(input_path: Path, output_pdf: Path) -> Optional[str]:

    if gencache is None:

        return "pywin32 が未インストールのため Office→PDF 変換できません。pip install pywin32 を実行してください。"

 

    suffix = input_path.suffix.lower()

    if suffix in {".doc", ".docx"}:

        word = None

        doc = None

        try:

            word = gencache.EnsureDispatch("Word.Application")

            word.Visible = False

            doc = word.Documents.Open(str(input_path), ReadOnly=True)

            doc.ExportAsFixedFormat(str(output_pdf), 17)

            return None

        except Exception as exc:

            return f"Word変換失敗: {exc}"

        finally:

            if doc is not None:

                try:

                    doc.Close(False)

                except Exception:

                    pass

            if word is not None:

                try:

                    word.Quit()

                except Exception:

                    pass

 

    if suffix in {".xls", ".xlsx", ".xlsm"}:

        excel = None

        wb = None

        try:

            excel = gencache.EnsureDispatch("Excel.Application")

            excel.Visible = False

            excel.DisplayAlerts = False

            wb = excel.Workbooks.Open(str(input_path))

            wb.ExportAsFixedFormat(0, str(output_pdf))

            return None

        except Exception as exc:

            return f"Excel変換失敗: {exc}"

        finally:

            if wb is not None:

                try:

                    wb.Close(False)

                except Exception:

                    pass

            if excel is not None:

                try:

                    excel.Quit()

                except Exception:

                    pass

 

    return f"未対応拡張子です: {suffix}"

 

def render_pdf_to_pngs(pdf_path: Path, image_dir: Path) -> Tuple[List[Path], Optional[str]]:

    if fitz is None:

        return [], "PyMuPDF が未インストールのため PDF→PNG 変換できません。pip install pymupdf を実行してください。"

 

    image_dir.mkdir(parents=True, exist_ok=True)

    try:

        doc = fitz.open(str(pdf_path))

    except Exception as exc:

        return [], f"PDF画像化失敗: {exc}"

 

    out_files: List[Path] = []

    try:

        for page_no, page in enumerate(doc, start=1):

            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)

            out = image_dir / f"page_{page_no:03d}.png"

            pix.save(str(out))

            out_files.append(out)

    except Exception as exc:

        return out_files, f"ページ画像化中に失敗: {exc}"

    finally:

        doc.close()

 

    return out_files, None

 

def sanitize_filename_for_path(name: str) -> str:

    cleaned = re.sub(r"[^A-Za-z0-9_.-]+", "_", name or "sheet")

    return cleaned[:40] or "sheet"

 

def is_excel_no_print_target_error(exc: Exception) -> bool:

    msg = str(exc)

    return ("印刷する対象がありません" in msg) or ("-2146827284" in msg)

 

def is_excel_sheet_export_skippable_error(exc: Exception) -> bool:

    """Treat common COM sheet-export failures as non-fatal and continue with other sheets."""

    msg = str(exc)

    return (

        is_excel_no_print_target_error(exc)

        or ("-2147024809" in msg)

        or ("無効な引数" in msg)

        or ("Invalid argument" in msg)

    )

 

def summarize_excel_sheet_export_error(exc: Exception) -> str:

    """Return a short, log-friendly reason for a skipped sheet export."""

    msg = str(exc)

    if "-2147024809" in msg or "無効な引数" in msg or "Invalid argument" in msg:

        return "COM無効引数"

    if is_excel_no_print_target_error(exc):

        return "印刷対象なし"

    return "COM例外"

 

def append_skipped_sheet(skipped_sheets: List[str], sheet_name: str, exc: Optional[Exception] = None) -> None:

    """Best-effort append that never raises during error handling."""

    reason = "COM例外"

    if exc is not None:

        try:

            reason = summarize_excel_sheet_export_error(exc)

        except Exception:

            reason = "COM例外"

    try:

        skipped_sheets.append(f"{sheet_name}({reason})")

    except Exception:

        skipped_sheets.append(str(sheet_name))

 

def append_sheet(skipped_sheets: List[str], sheet_name: str, exc: Optional[Exception] = None) -> None:

    """Backward-compatible alias for older call sites."""

    append_skipped_sheet(skipped_sheets, sheet_name, exc)

 

def convert_excel_to_sheet_pdfs(input_path: Path, pdf_dir: Path, slug: str) -> Tuple[List[Tuple[str, Path]], List[str], Optional[str]]:

    if gencache is None:

        return [], [], "pywin32 が未インストールのため Excelシート別PDF変換できません。pip install pywin32 を実行してください。"

 

    excel = None

    wb = None

    out: List[Tuple[str, Path]] = []

    skipped_sheets: List[str] = []

    try:

        excel = gencache.EnsureDispatch("Excel.Application")

        excel.Visible = False

        excel.DisplayAlerts = False

        wb = excel.Workbooks.Open(str(input_path), ReadOnly=True)

 

        sheet_count = int(getattr(wb.Worksheets, "Count", 0) or 0)

        for idx in range(1, sheet_count + 1):

            sheet_name = f"Sheet{idx}"

            try:

                ws = wb.Worksheets(idx)

                sheet_name = str(getattr(ws, "Name", sheet_name))

            except Exception as exc:

                append_skipped_sheet(skipped_sheets, sheet_name, exc)

                continue

            safe_sheet = sanitize_filename_for_path(sheet_name)

            pdf_path = pdf_dir / f"{slug}_s{idx:03d}_{safe_sheet}.pdf"

            if pdf_path.exists():

                try:

                    pdf_path.unlink()

                except Exception:

                    pass

            try:

                ws.ExportAsFixedFormat(0, str(pdf_path))

            except Exception as exc:

                try:

                    # Retry once by ignoring broken print areas.

                    ws.ExportAsFixedFormat(0, str(pdf_path), IgnorePrintAreas=True)

                except Exception as retry_exc:

                    append_skipped_sheet(skipped_sheets, sheet_name, retry_exc)

                    continue

            if pdf_path.exists():

                out.append((sheet_name, pdf_path))

            else:

                skipped_sheets.append(sheet_name)

 

        if not out:

            # Fallback: workbook whole export (some layouts cannot be exported sheet-by-sheet).

            fallback_pdf = pdf_dir / f"{slug}_workbook.pdf"

            try:

                wb.ExportAsFixedFormat(0, str(fallback_pdf))

                if fallback_pdf.exists():

                    return [("(Workbook)", fallback_pdf)], skipped_sheets, None

            except Exception as exc:

                try:

                    wb.ExportAsFixedFormat(0, str(fallback_pdf), IgnorePrintAreas=True)

                    if fallback_pdf.exists():

                        return [("(Workbook)", fallback_pdf)], skipped_sheets, None

                except Exception as retry_exc:

                    return [], skipped_sheets, f"Excelシート別PDF変換に失敗しました（出力0件）: {retry_exc}"

            return [], skipped_sheets, "Excelシート別PDF変換に失敗しました（出力0件）。"

        return out, skipped_sheets, None

    except Exception as exc:

        return out, skipped_sheets, f"Excelシート別変換失敗: {exc}"

    finally:

        if wb is not None:

            try:

                wb.Close(False)

            except Exception:

                pass

        if excel is not None:

            try:

                excel.Quit()

            except Exception:

                pass

 

def run_visual_pipeline(

    file_path: Path,

    results: List[CheckResult],

    visual_pages: List[VisualPage],

    assets_root: Path,

) -> Optional[int]:

    suffix = file_path.suffix.lower()

    slug = slugify_for_path(file_path)

    pdf_dir = assets_root / "pdf"

    image_dir = assets_root / "png" / slug

    pdf_dir.mkdir(parents=True, exist_ok=True)

 

    if suffix in {".xls", ".xlsx", ".xlsm"}:

        sheet_pdfs, skipped_sheets, sheet_err = convert_excel_to_sheet_pdfs(file_path, pdf_dir, slug)

        if sheet_err:

            add_result(

                results,

                file_path,

                "CommonVisual",

                "V1",

                "共通PDF出力・PNG化",

                "ERROR",

                sheet_err,

                "OfficeアプリがインストールされたWindows環境で再実行してください。",

            )

            return None

 

        if skipped_sheets:

            add_result(

                results,

                file_path,

                "CommonVisual",

                "V1-WARN",

                "共通PDF出力・PNG化（シート別）",

                "WARN",

                f"印刷対象なし等でシート別PDF化をスキップ: {', '.join(skipped_sheets[:6])}"

                + (f" ほか{len(skipped_sheets)-6}件" if len(skipped_sheets) > 6 else ""),

                "対象シートの印刷範囲・印刷設定を確認してください。",

            )

 

        total_pages = 0

        global_page_no = 1

        for sheet_name, sheet_pdf in sheet_pdfs:

            sheet_image_dir = image_dir / sanitize_filename_for_path(sheet_name)

            images, image_err = render_pdf_to_pngs(sheet_pdf, sheet_image_dir)

            if image_err:

                add_result(

                    results,

                    file_path,

                    "CommonVisual",

                    "V1",

                    "共通PDF出力・PNG化",

                    "ERROR",

                    image_err,

                    "必要ライブラリを導入し、再実行してください。",

                )

                return None

            for image_path in images:

                visual_pages.append(

                    VisualPage(str(file_path), global_page_no, str(image_path), sheet_name=sheet_name)

                )

                global_page_no += 1

                total_pages += 1

 

        add_result(

            results,

            file_path,

            "CommonVisual",

            "V1",

            "共通PDF出力・PNG化",

            "PASS",

            f"{total_pages}ページをPNG化し、別シート出力用データを作成。",

            "画像シートで見切れ/罫線欠け/表サイズ/ページ番号を確認。",

        )

        return total_pages

 

    pdf_path: Optional[Path] = None

    if suffix == ".pdf":

        pdf_path = file_path

    elif suffix in {".doc", ".docx"}:

        pdf_path = pdf_dir / f"{slug}.pdf"

        err = convert_office_to_pdf(file_path, pdf_path)

        if err:

            add_result(

                results,

                file_path,

                "CommonVisual",

                "V1",

                "共通PDF出力・PNG化",

                "ERROR",

                err,

                "OfficeアプリがインストールされたWindows環境で再実行してください。",

            )

            return None

    else:

        add_result(

            results,

            file_path,

            "CommonVisual",

            "V1",

            "共通PDF出力・PNG化",

            "N/A",

            "この拡張子は共通PDF出力対象外です。",

            "対応不要。",

        )

        return None

 

    if pdf_path is None or (not pdf_path.exists()):

        add_result(

            results,

            file_path,

            "CommonVisual",

            "V1",

            "共通PDF出力・PNG化",

            "ERROR",

            "PDF生成に失敗しました。",

            "変換環境を確認して再実行してください。",

        )

        return None

 

    images, image_err = render_pdf_to_pngs(pdf_path, image_dir)

    if image_err:

        add_result(

            results,

            file_path,

            "CommonVisual",

            "V1",

            "共通PDF出力・PNG化",

            "ERROR",

            image_err,

            "必要ライブラリを導入し、再実行してください。",

        )

        return None

 

    for idx, image_path in enumerate(images, start=1):

        visual_pages.append(VisualPage(str(file_path), idx, str(image_path), sheet_name=""))

 

    add_result(

        results,

        file_path,

        "CommonVisual",

        "V1",

        "共通PDF出力・PNG化",

        "PASS",

        f"{len(images)}ページをPNG化し、別シート出力用データを作成。",

        "画像シートで見切れ/罫線欠け/表サイズ/ページ番号を確認。",

    )

    return len(images)

 

def append_max_page_detail(results: List[CheckResult], file_path: Path, page_count: Optional[int]) -> None:

    if page_count is None:

        return

    target = str(file_path)

    for r in results:

        if r.file_path != target:

            continue

        if "最大ページ数=" in r.detail:

            continue

        r.detail = f"{r.detail} / 最大ページ数={page_count}"

        if (

            r.file_type == "Word"

            and r.status in {"FAIL", "WARN", "MANUAL"}

            and "該当ページ=" not in r.detail

        ):

            r.detail = f"{r.detail} / 該当ページ=自動特定不可(Wordレイアウト依存)"

 

def derive_applicability(status: str) -> str:

    if status == "N/A":

        return "NOT_APPLICABLE"

    return "APPLICABLE"

 

def derive_automation(status: str) -> str:

    if status == "N/A":

        return "NOT_SUPPORTED"

    if status == "MANUAL":

        return "MANUAL_ONLY"

    return "AUTO_CHECKED"

 

def display_status(status: str) -> str:

    mapping = {

        "FAIL": "×",

        "WARN": "警告(要確認)",

        "MANUAL": "目検で確認",

        "N/A": "対象外",

        "PASS": "〇",

    }

    return mapping.get(status, status)

 

def populate_image_preview_sheet(

    ws_images,

    visual_pages: Iterable[VisualPage],

    preview_pages_per_row: int = 6,

) -> None:

    ws_images.append(["file_path", "page_count", "layout"])

    ws_images.column_dimensions["A"].width = 65

    ws_images.column_dimensions["B"].width = 12

    ws_images.column_dimensions["C"].width = 28

    ws_images.freeze_panes = "D2"

 

    unique_visual_pages = {}

    for page in visual_pages:

        key = (page.file_path, page.page_no)

        if key in unique_visual_pages:

            continue

        unique_visual_pages[key] = page

 

    pages_by_file = defaultdict(list)

    for page in unique_visual_pages.values():

        pages_by_file[page.file_path].append(page)

 

    for file_path in pages_by_file:

        pages_by_file[file_path] = sorted(pages_by_file[file_path], key=lambda p: p.page_no)

 

    _ = preview_pages_per_row  # 後方互換のため引数は保持（レイアウトは横スクロール固定）

    thumb_width = 260

    default_image_row_height = 210

    block_gap_rows = 3

    start_col = 4

    col_span_per_page = 3

 

    cursor_row = 2

    for file_path in sorted(pages_by_file.keys()):

        pages = pages_by_file[file_path]

        ws_images.cell(row=cursor_row, column=1, value=file_path)

        ws_images.cell(row=cursor_row, column=2, value=len(pages))

        ws_images.cell(row=cursor_row, column=3, value="横スクロールで全ページ確認")

 

        if not pages:

            cursor_row += 3

            continue

 

        base_row = cursor_row + 1

        sheet_row = cursor_row + 2

        max_image_height_pt = 0.0

        for idx, page in enumerate(pages):

            base_col = start_col + (idx * col_span_per_page)

            for col in range(base_col, base_col + col_span_per_page):

                col_letter = ws_images.cell(row=1, column=col).column_letter

                ws_images.column_dimensions[col_letter].width = 22

 

            ws_images.cell(row=cursor_row, column=base_col, value=f"p{page.page_no}")

            ws_images.cell(

                row=sheet_row,

                column=base_col,

                value=f"シート名：{page.sheet_name}" if page.sheet_name else "",

            )

 

            if XLImage is not None and Path(page.image_path).exists():

                try:

                    img = XLImage(page.image_path)

                    if getattr(img, "width", 0) and img.width > thumb_width:

                        ratio = thumb_width / img.width

                        img.width = int(img.width * ratio)

                        img.height = int(img.height * ratio)

                    img_height_pt = (float(getattr(img, "height", 0) or 0) / PIXELS_PER_INCH) * POINTS_PER_INCH

                    if img_height_pt > 0:

                        max_image_height_pt = max(max_image_height_pt, img_height_pt)

                    ws_images.add_image(img, ws_images.cell(row=base_row, column=base_col + 1).coordinate)

                except Exception as exc:

                    ws_images.cell(row=base_row, column=base_col + 1, value=f"画像貼付失敗: {exc}")

            else:

                ws_images.cell(row=base_row, column=base_col + 1, value="画像ファイルなし")

 

        if max_image_height_pt > 0:

            ws_images.row_dimensions[base_row].height = max_image_height_pt + mm_to_points(2.0)

        else:

            ws_images.row_dimensions[base_row].height = default_image_row_height

        ws_images.row_dimensions[sheet_row].height = 18

        cursor_row = cursor_row + 3 + block_gap_rows

 

def write_results_report_xlsx(results: Iterable[CheckResult], output_xlsx: Path) -> Optional[str]:

    if Workbook is None:

        return "openpyxl が未インストールのため結果xlsxを出力できません。"

 

    output_xlsx.parent.mkdir(parents=True, exist_ok=True)

    wb = Workbook()

    ws_results = wb.active

    ws_results.title = "results"

    ws_results.append(

        [

            "file_path",

            "file_type",

            "check_id",

            "check_item",

            "applicability",

            "automation",

            "status",

            "detail",

            "suggested_action",

        ]

    )

    for r in results:

        ws_results.append(

            [

                r.file_path,

                r.file_type,

                r.check_id,

                r.check_item,

                derive_applicability(r.status),

                derive_automation(r.status),

                display_status(r.status),

                r.detail,

                r.suggested_action,

            ]

        )

 

    try:

        wb.save(output_xlsx)

        return None

    except Exception as exc:

        return f"results xlsx出力失敗: {exc}"

 

def write_image_preview_xlsx(

    visual_pages: Iterable[VisualPage],

    output_xlsx: Path,

    preview_pages_per_row: int = 6,

) -> Optional[str]:

    if Workbook is None:

        return "openpyxl が未インストールのためimage_preview xlsxを出力できません。"

 

    output_xlsx.parent.mkdir(parents=True, exist_ok=True)

    wb = Workbook()

    ws_images = wb.active

    ws_images.title = "image_preview"

    populate_image_preview_sheet(ws_images, visual_pages, preview_pages_per_row=preview_pages_per_row)

 

    try:

        wb.save(output_xlsx)

        return None

    except Exception as exc:

        return f"image_preview xlsx出力失敗: {exc}"

 

def write_visual_report_xlsx(

    results: Iterable[CheckResult],

    visual_pages: Iterable[VisualPage],

    output_xlsx: Path,

    preview_pages_per_row: int = 6,

) -> Optional[str]:

    if Workbook is None:

        return "openpyxl が未インストールのため画像シートxlsxを出力できません。"

 

    output_xlsx.parent.mkdir(parents=True, exist_ok=True)

    wb = Workbook()

    ws_results = wb.active

    ws_results.title = "results"

    ws_results.append(

        [

            "file_path",

            "file_type",

            "check_id",

            "check_item",

            "applicability",

            "automation",

            "status",

            "detail",

            "suggested_action",

        ]

    )

    for r in results:

        ws_results.append(

            [

                r.file_path,

                r.file_type,

                r.check_id,

                r.check_item,

                derive_applicability(r.status),

                derive_automation(r.status),

                display_status(r.status),

                r.detail,

                r.suggested_action,

            ]

        )

 

    ws_images = wb.create_sheet("image_preview")

    populate_image_preview_sheet(ws_images, visual_pages, preview_pages_per_row=preview_pages_per_row)

 

    try:

        wb.save(output_xlsx)

        return None

    except Exception as exc:

        return f"xlsx出力失敗: {exc}"

 

def coord_in_ranges(coord: str, ranges: List[Tuple[int, int, int, int]]) -> bool:

    row, col = coordinate_to_tuple(coord)

    for min_col, min_row, max_col, max_row in ranges:

        if min_row <= row <= max_row and min_col <= col <= max_col:

            return True

    return False

 

def collect_formula_refs_excel(workbook) -> Set[str]:

    refs: Set[str] = set()

    for ws in workbook.worksheets:

        cells = getattr(ws, "_cells", None)

        if isinstance(cells, dict):

            iter_cells = cells.values()

        else:

            iter_cells = (cell for _, _, cell in iter_nonempty_cells(ws))

        for cell in iter_cells:

            value = getattr(cell, "value", None)

            if isinstance(value, str) and value.startswith("="):

                refs.update(extract_cell_refs(value))

    return refs

 

def extract_cell_refs(formula: str) -> Set[str]:

    refs = set(re.findall(r"\b([A-Za-z]{1,3}\d{1,7})\b", formula))

    return {r.upper() for r in refs}

 

def parse_print_area_ranges(ws) -> List[Tuple[int, int, int, int]]:

    areas: List[str] = []

    if ws.print_area:

        if isinstance(ws.print_area, str):

            areas = [ws.print_area]

        else:

            areas = list(ws.print_area)

    parsed: List[Tuple[int, int, int, int]] = []

    for area in areas:

        cleaned = area.replace("$", "")

        if "!" in cleaned:

            cleaned = cleaned.split("!", 1)[1]

        for part in cleaned.split(","):

            part = part.strip()

            if ":" not in part:

                continue

            try:

                parsed.append(range_boundaries(part))

            except Exception:

                continue

    return parsed

 

def iter_nonempty_cells(ws):

    """Yield only cells that actually have values to avoid scanning huge empty ranges."""

    cells = getattr(ws, "_cells", None)

    if isinstance(cells, dict):

        for (row_idx, col_idx), cell in sorted(cells.items()):

            if cell is None or cell.value is None:

                continue

            yield row_idx, col_idx, cell

        return

 

    max_row = ws.max_row or 1

    max_col = ws.max_column or 1

    for row_idx in range(1, max_row + 1):

        for col_idx in range(1, max_col + 1):

            cell = ws.cell(row=row_idx, column=col_idx)

            if cell.value is None:

                continue

            yield row_idx, col_idx, cell

 

def extract_excel_break_ids(break_container) -> List[int]:

    if break_container is None:

        return []

    ids: List[int] = []

    for br in getattr(break_container, "brk", []) or []:

        try:

            bid = int(getattr(br, "id"))

            ids.append(bid)

        except Exception:

            continue

    return sorted(set(ids))

 

def infer_excel_print_page(ws, row_idx: int, col_idx: int) -> str:

    row_break_ids = extract_excel_break_ids(getattr(ws, "row_breaks", None))

    col_break_ids = extract_excel_break_ids(getattr(ws, "col_breaks", None))

    return infer_excel_print_page_from_breaks(row_break_ids, col_break_ids, row_idx, col_idx)

 

def infer_excel_print_page_from_breaks(

    row_break_ids: List[int],

    col_break_ids: List[int],

    row_idx: int,

    col_idx: int,

) -> str:

    row_page = 1 + bisect_right(row_break_ids, row_idx - 1)

    col_page = 1 + bisect_right(col_break_ids, col_idx - 1)

    if col_page == 1:

        return f"P{row_page}"

    return f"P{row_page}-{col_page}"

 

def summarize_locations(locations: List[str], limit: int = 8) -> str:

    if not locations:

        return ""

    shown = locations[:limit]

    rest = len(locations) - len(shown)

    if rest > 0:

        return ", ".join(shown) + f" ほか{rest}件"

    return ", ".join(shown)

 

def summarize_pages(page_labels: List[str], limit: int = 8) -> str:

    if not page_labels:

        return "指摘対象ページ：なし"

    ordered = list(dict.fromkeys(page_labels))

    shown = ordered[:limit]

    rest = len(ordered) - len(shown)

    if rest > 0:

        return "指摘対象ページ：" + ", ".join(shown) + f" ほか{rest}件"

    return "指摘対象ページ：" + ", ".join(shown)

 

def check_excel(file_path: Path, results: List[CheckResult]) -> None:

    if load_workbook is None:

        add_result(

            results,

            file_path,

            "Excel",

            "E-ENV",

            "Excel解析ライブラリ",

            "ERROR",

            "openpyxl が未インストールのため解析できません。",

            "pip install openpyxl を実行してください。",

        )

        return

 

    try:

        with warnings.catch_warnings():

            warnings.filterwarnings(

                "ignore",

                message=r"(?i)cannot parse header or footer so it will be ignored",

                category=UserWarning,

            )

            wb = load_workbook(file_path, data_only=False)

    except Exception as exc:

        add_result(

            results,

            file_path,

            "Excel",

            "E-OPEN",

            "Excelファイル読込",

            "ERROR",

            f"読込失敗: {exc}",

            "ファイル破損・パスワード保護の有無を確認してください。",

        )

        return

 

    all_formula_refs = collect_formula_refs_excel(wb)

 

    strike_cells = 0

    comment_cells = 0

    strike_locations: List[str] = []

    comment_locations: List[str] = []

    ref_error_locations: List[str] = []

    strike_pages: List[str] = []

    comment_pages: List[str] = []

    ref_error_pages: List[str] = []

    missing_print_title_rows_sheets: List[str] = []

 

    ref_error_tokens = {

        "#REF!",

        "#NAME?",

        "#VALUE!",

        "#DIV/0!",

        "#NUM!",

        "#NULL!",

        "#N/A",

    }

 

    for ws in wb.worksheets:

        print_title_rows = getattr(ws, "print_title_rows", None)

        if isinstance(print_title_rows, str) and print_title_rows.strip():

            pass

        else:

            missing_print_title_rows_sheets.append(ws.title)

 

        print_ranges = parse_print_area_ranges(ws)

        row_break_ids = extract_excel_break_ids(getattr(ws, "row_breaks", None))

        col_break_ids = extract_excel_break_ids(getattr(ws, "col_breaks", None))

        outside_cells: List[str] = []

        outside_locations: List[str] = []

        outside_referenced_locations: List[str] = []

        outside_pages: List[str] = []

        outside_referenced_pages: List[str] = []

 

        for row_idx, col_idx, cell in iter_nonempty_cells(ws):

            page_label = infer_excel_print_page_from_breaks(row_break_ids, col_break_ids, row_idx, col_idx)

            coord = f"{get_column_letter(col_idx)}{row_idx}"

            loc = f"{ws.title}!{coord}({page_label})"

 

            if bool(getattr(getattr(cell, "font", None), "strike", False)):

                strike_cells += 1

                strike_locations.append(loc)

                strike_pages.append(page_label)

            if getattr(cell, "comment", None) is not None:

                comment_cells += 1

                comment_locations.append(loc)

                comment_pages.append(page_label)

 

            cell_text = str(cell.value).upper() if isinstance(cell.value, str) else ""

            if (isinstance(cell.value, str) and cell.value.startswith("=") and "#REF!" in cell_text) or (

                cell_text in ref_error_tokens

            ):

                ref_error_locations.append(loc)

                ref_error_pages.append(page_label)

 

            if print_ranges and (not coord_in_ranges(coord, print_ranges)):

                outside_cells.append(coord)

                outside_locations.append(loc)

                outside_pages.append(page_label)

                if coord in all_formula_refs:

                    outside_referenced_locations.append(loc)

                    outside_referenced_pages.append(page_label)

 

        if not print_ranges:

            add_result(

                results,

                file_path,

                "Excel",

                "E1",

                f"印刷範囲外記載チェック [{ws.title}]",

                "WARN",

                "印刷範囲が未設定のため、自動判定できません。 / 指摘対象ページ：特定不可(印刷範囲未設定)",

                "必要なら印刷範囲を設定し、再チェックしてください。",

            )

            continue

 

        if outside_cells:

            add_result(

                results,

                file_path,

                "Excel",

                "E1",

                f"印刷範囲外記載チェック [{ws.title}]",

                "FAIL",

                (

                    f"印刷範囲外セル {len(outside_cells)} 件（うち参照あり {len(outside_referenced_locations)} 件）。"

                    f" / {summarize_pages(outside_pages)}"

                    f" / 該当セル: {summarize_locations(outside_locations)}"

                    + (

                        f" / 参照ありセル: {summarize_locations(outside_referenced_locations)} / {summarize_pages(outside_referenced_pages)}"

                        if outside_referenced_locations

                        else ""

                    )

                ),

                "必要なら印刷範囲拡大、不要かつ未参照なら削除、参照ありは残置を検討。",

            )

        else:

            add_result(

                results,

                file_path,

                "Excel",

                "E1",

                f"印刷範囲外記載チェック [{ws.title}]",

                "PASS",

                "印刷範囲外の記載は検出されませんでした。 / 指摘対象ページ：なし",

                "対応不要。",

            )

 

    all_sheet_names = [ws.title for ws in wb.worksheets]

    add_result(

        results,

        file_path,

        "Excel",

        "E2",

        "不要シートチェック",

        "MANUAL",

        f"対象シート: {', '.join(all_sheet_names)}",

        "不要シート有無は対象シート一覧をもとに手動確認。",

    )

 

    add_result(

        results,

        file_path,

        "Excel",

        "E3",

        "印刷プレビュー（見切れ/罫線欠け）",

        "MANUAL",

        "プレビュー見切れ・罫線欠けは自動判定が難しいため目視確認が必要です。",

        "Excelで印刷プレビューを開き、見切れや罫線欠けを補正。",

    )

 

    add_result(

        results,

        file_path,

        "Excel",

        "E4",

        "見え消し（取り消し線/訂正線）残存",

        "FAIL" if strike_cells > 0 else "PASS",

        (

            f"取り消し線セル: {strike_cells} 件。 / {summarize_pages(strike_pages)} / 該当セル: {summarize_locations(strike_locations)}"

            if strike_cells > 0

            else "取り消し線セル: 0 件。 / 指摘対象ページ：なし"

        ),

        "不要な取り消し線を解除。" if strike_cells > 0 else "対応不要。",

    )

 

    add_result(

        results,

        file_path,

        "Excel",

        "E5",

        "コメント（吹き出し）残存",

        "FAIL" if comment_cells > 0 else "PASS",

        (

            f"コメント付きセル: {comment_cells} 件。 / {summarize_pages(comment_pages)} / 該当セル: {summarize_locations(comment_locations)}"

            if comment_cells > 0

            else "コメント付きセル: 0 件。 / 指摘対象ページ：なし"

        ),

        "不要なコメント/メモを削除。" if comment_cells > 0 else "対応不要。",

    )

 

    add_result(

        results,

        file_path,

        "Excel",

        "E6",

        "参照エラー残存",

        "FAIL" if ref_error_locations else "PASS",

        (

            f"参照エラー/数式エラーセル: {len(ref_error_locations)} 件。 / {summarize_pages(ref_error_pages)} / 該当セル: {summarize_locations(ref_error_locations)}"

            if ref_error_locations

            else "参照エラー/数式エラーセルは検出されません。 / 指摘対象ページ：なし"

        ),

        "該当セルの数式参照先を修正。" if ref_error_locations else "対応不要。",

    )

 

    add_result(

        results,

        file_path,

        "Excel",

        "E7",

        "ページ設定のタイトル行",

        "FAIL" if missing_print_title_rows_sheets else "PASS",

        (

            "タイトル行が未設定のシートがあります: " + ", ".join(missing_print_title_rows_sheets)

            if missing_print_title_rows_sheets

            else "全シートでタイトル行が設定されています。"

        ),

        "ページ設定のタイトル行（先頭行）を設定してください。"

        if missing_print_title_rows_sheets

        else "対応不要。",

    )

 

    cp = wb.properties

    has_props = any(

        [

            cp.creator,

            cp.lastModifiedBy,

            cp.title,

            cp.subject,

            cp.keywords,

            cp.description,

            cp.category,

            cp.contentStatus,

        ]

    )

    add_result(

        results,

        file_path,

        "Excel",

        "C1",

        "プロパティ情報削除",

        "FAIL" if has_props else "PASS",

        "プロパティ情報が残っています。" if has_props else "主要プロパティは空です。",

        "ファイル情報のプロパティを削除。" if has_props else "対応不要。",

    )

 

    add_result(

        results,

        file_path,

        "Excel",

        "C2",

        "表紙が規定のもの",

        "MANUAL",

        "規定表紙判定は基準データが必要なため手動確認です。",

        "規定表紙テンプレートと照合。",

    )

 

    margin_msgs = []

    margin_fail = False

    for ws in wb.worksheets:

        pm = ws.page_margins

        top = inches_to_mm(getattr(pm, "top", None))

        bottom = inches_to_mm(getattr(pm, "bottom", None))

        left = inches_to_mm(getattr(pm, "left", None))

        right = inches_to_mm(getattr(pm, "right", None))

        if None in (top, bottom, left, right):

            margin_msgs.append(f"{ws.title}(余白取得不可)")

            margin_fail = True

            continue

        margin_msgs.append(f"{ws.title}(上{top:.1f}/下{bottom:.1f}/左{left:.1f}/右{right:.1f}mm)")

        if not (top >= 20 and bottom >= 20 and left >= 30 and right >= 20):

            margin_fail = True

 

    add_result(

        results,

        file_path,

        "Excel",

        "C3",

        COMMON_MARGIN_RULE_TEXT,

        "FAIL" if margin_fail else "PASS",

        "; ".join(margin_msgs),

        "ページ設定で余白を補正。" if margin_fail else "設定値確認済み。",

    )

 

    add_result(

        results,

        file_path,

        "Excel",

        "C4",

        "ページ番号",

        "MANUAL",

        "Excelヘッダ/フッタのページ番号形式は目視確認を推奨。最大ページ数は画像シート参照。",

        "印刷プレビューまたはページ設定で確認。",

    )

 

    add_result(

        results,

        file_path,

        "Excel",

        "C5",

        "PDF出力結果確認（見切れ/罫線/表サイズ/ページ番号）",

        "MANUAL",

        "PDF出力後の品質確認は目視が必要です。",

        "PDFを開き、見切れ・罫線欠け・表サイズ・ページ番号を確認。",

    )

 

    try:

        wb.close()

    except Exception:

        pass

 

def mm_from_emu(value) -> Optional[float]:

    if value is None:

        return None

    try:

        return float(value) / EMU_PER_MM

    except Exception:

        return None

 

def normalize_word_snippet(text: str, max_len: int = 60) -> str:

    cleaned = re.sub(r"\s+", " ", (text or "")).strip()

    return cleaned[:max_len]

 

def find_word_text_page_numbers(

    doc_path: Path,

    snippets: List[str],

    max_snippets: int = 30,

    max_hits_per_snippet: int = 5,

    max_seconds: float = 8.0,

) -> Dict[str, List[int]]:

    if gencache is None:

        return {}

 

    uniq_snippets: List[str] = []

    seen = set()

    for snippet in snippets:

        s = normalize_word_snippet(snippet)

        if (not s) or (len(s) < 3) or (s in seen):

            continue

        seen.add(s)

        uniq_snippets.append(s)

    if not uniq_snippets:

        return {}

 

    word = None

    doc = None

    page_map: Dict[str, List[int]] = {}

    try:

        word = gencache.EnsureDispatch("Word.Application")

        word.Visible = False

        doc = word.Documents.Open(str(doc_path), ReadOnly=True)

        content_end = doc.Content.End

        search_start = time.perf_counter()

 

        for snippet in uniq_snippets[:max_snippets]:

            if (time.perf_counter() - search_start) >= max_seconds:

                break

            rng = doc.Content

            pages = set()

            steps = 0

            last_end = -1

            while True:

                if (time.perf_counter() - search_start) >= max_seconds:

                    break

                if steps >= 200:

                    break

                try:

                    finder = rng.Find

                    finder.ClearFormatting()

                    finder.Text = snippet

                    finder.Forward = True

                    finder.Wrap = 0

                    finder.MatchWildcards = False

                    finder.MatchWholeWord = False

                    finder.MatchCase = False

                    executed = bool(finder.Execute())

                except Exception:

                    break

                if not executed:

                    break

                try:

                    page_no = int(rng.Information(3))

                except Exception:

                    page_no = None

                if page_no:

                    pages.add(page_no)

                if len(pages) >= max_hits_per_snippet:

                    break

                next_start = rng.End

                if next_start >= content_end:

                    break

                if next_start <= last_end:

                    break

                try:

                    rng.SetRange(next_start, content_end)

                except Exception:

                    break

                last_end = next_start

                steps += 1

 

            if pages:

                page_map[snippet] = sorted(pages)

    except Exception:

        return {}

    finally:

        if doc is not None:

            try:

                doc.Close(False)

            except Exception:

                pass

        if word is not None:

            try:

                word.Quit()

            except Exception:

                pass

 

    return page_map

 

def collect_pages_from_snippets(snippets: List[str], page_map: Dict[str, List[int]]) -> List[str]:

    pages = set()

    for snippet in snippets:

        s = normalize_word_snippet(snippet)

        for page_no in page_map.get(s, []):

            pages.add(f"P{page_no}")

    return sorted(pages)

 

def iter_word_runs(doc):

    for paragraph in doc.paragraphs:

        for run in paragraph.runs:

            yield run

    for table in doc.tables:

        for row in table.rows:

            for cell in row.cells:

                for paragraph in cell.paragraphs:

                    for run in paragraph.runs:

                        yield run

 

def check_word(file_path: Path, results: List[CheckResult], cover_keyword: Optional[str]) -> None:

    if Document is None:

        add_result(

            results,

            file_path,

            "Word",

            "W-ENV",

            "Word解析ライブラリ",

            "ERROR",

            "python-docx が未インストールのため解析できません。",

            "pip install python-docx を実行してください。",

        )

        return

 

    try:

        doc = Document(file_path)

    except Exception as exc:

        add_result(

            results,

            file_path,

            "Word",

            "W-OPEN",

            "Wordファイル読込",

            "ERROR",

            f"読込失敗: {exc}",

            "ファイル破損・保護設定を確認してください。",

        )

        return

 

    doc_xml = doc.element.xml

    settings_xml = doc.settings.element.xml if getattr(doc, "settings", None) is not None else ""

    revision_patterns = [r"<w:ins\\b", r"<w:del\\b", r"<w:moveFrom\\b", r"<w:moveTo\\b"]

    revision_count = sum(len(re.findall(p, doc_xml)) for p in revision_patterns)

    track_revisions_on = "w:trackRevisions" in settings_xml

    strike_count = 0

    double_strike_count = 0

    highlighted_runs = 0

    colored_runs = 0

    strike_snippets: List[str] = []

    highlight_snippets: List[str] = []

    colored_snippets: List[str] = []

    ref_error_snippets: List[str] = []

    for run in iter_word_runs(doc):

        run_text = normalize_word_snippet(getattr(run, "text", ""))

        if bool(getattr(run.font, "strike", False)):

            strike_count += 1

            if run_text:

                strike_snippets.append(run_text)

        if bool(getattr(run.font, "double_strike", False)):

            double_strike_count += 1

            if run_text:

                strike_snippets.append(run_text)

        if run.font.highlight_color is not None:

            highlighted_runs += 1

            if run_text:

                highlight_snippets.append(run_text)

        if run.font.color is not None and run.font.color.rgb is not None:

            colored_runs += 1

            if run_text:

                colored_snippets.append(run_text)

 

    has_track_changes = revision_count > 0 or track_revisions_on

    has_miekeshi = has_track_changes or (strike_count + double_strike_count) > 0

    add_result(

        results,

        file_path,

        "Word",

        "W1",

        "見え消し（変更履歴）残存",

        "FAIL" if has_miekeshi else "PASS",

        (

            f"変更履歴要素={revision_count}, trackRevisions={track_revisions_on}, 取り消し線Run={strike_count}, 二重取り消し線Run={double_strike_count}。 / 指摘対象ページ：特定不可(Wordレイアウト依存)"

            if has_miekeshi

            else "見え消し（変更履歴/取り消し線/二重取り消し線）は検出されません。 / 指摘対象ページ：なし"

        ),

        "最新版確認後、変更履歴を承諾/破棄し、取り消し線を解消。" if has_miekeshi else "対応不要。",

    )

 

    add_result(

        results,

        file_path,

        "Word",

        "W2",

        "マーカ残存",

        "FAIL" if highlighted_runs > 0 else "PASS",

        (

            f"ハイライト付きRun: {highlighted_runs} 件。 / 指摘対象ページ：特定不可(Wordレイアウト依存)"

            if highlighted_runs > 0

            else "ハイライト付きRun: 0 件。 / 指摘対象ページ：なし"

        ),

        "最新版確認後、不要マーカを削除。" if highlighted_runs > 0 else "対応不要。",

    )

 

    add_result(

        results,

        file_path,

        "Word",

        "W3",

        "不要な文字色",

        "WARN" if colored_runs > 0 else "PASS",

        (

            f"明示的な文字色付きRun: {colored_runs} 件。 / 指摘対象ページ：特定不可(Wordレイアウト依存)"

            if colored_runs > 0

            else "明示的な文字色付きRun: 0 件。 / 指摘対象ページ：なし"

        ),

        "不要な色がないことを最終確認し、色付き箇所は標準色へ修正。",

    )

 

    comment_anchor_count = len(re.findall(r"<w:commentRangeStart\\b", doc_xml)) + len(

        re.findall(r"<w:commentReference\\b", doc_xml)

    )

    has_comment_part = any("/comments.xml" in str(p.partname) for p in doc.part.package.parts)

    has_comments = comment_anchor_count > 0 or has_comment_part

    add_result(

        results,

        file_path,

        "Word",

        "W4",

        "コメント（吹き出し）残存",

        "FAIL" if has_comments else "PASS",

        (

            f"コメント参照を検出（anchor={comment_anchor_count}, comments.xml={has_comment_part}）。 / 指摘対象ページ：特定不可(Wordレイアウト依存)"

            if has_comments

            else "コメント（吹き出し）は検出されません。 / 指摘対象ページ：なし"

        ),

        "不要なコメントを削除し、最終版で吹き出し非表示を確認。" if has_comments else "対応不要。",

    )

 

    word_text_blocks = [p.text for p in doc.paragraphs if p.text]

    for table in doc.tables:

        for row in table.rows:

            for cell in row.cells:

                for p in cell.paragraphs:

                    if p.text:

                        word_text_blocks.append(p.text)

    ref_error_pattern = re.compile(r"(?i)(error!\s*reference source not found\.?|#ref!)")

    word_ref_hits = 0

    for t in word_text_blocks:

        if ref_error_pattern.search(t):

            word_ref_hits += 1

            ref_error_snippets.append(normalize_word_snippet(t))

 

    all_lookup_snippets = strike_snippets + highlight_snippets + colored_snippets + ref_error_snippets

    word_page_map = find_word_text_page_numbers(file_path, all_lookup_snippets)

    w1_pages = collect_pages_from_snippets(strike_snippets, word_page_map)

    w2_pages = collect_pages_from_snippets(highlight_snippets, word_page_map)

    w3_pages = collect_pages_from_snippets(colored_snippets, word_page_map)

    w5_pages = collect_pages_from_snippets(ref_error_snippets, word_page_map)

    add_result(

        results,

        file_path,

        "Word",

        "W5",

        "参照エラー残存",

        "FAIL" if word_ref_hits > 0 else "PASS",

        (

            f"参照エラー文字列を {word_ref_hits} 件検出。 / 指摘対象ページ：特定不可(Wordレイアウト依存)"

            if word_ref_hits > 0

            else "参照エラー文字列は検出されません。 / 指摘対象ページ：なし"

        ),

        "参照元(相互参照・図表番号・ブックマーク)を再設定。" if word_ref_hits > 0 else "対応不要。",

    )

 

    cp = doc.core_properties

    has_props = any(

        [

            cp.author,

            cp.last_modified_by,

            cp.title,

            cp.subject,

            cp.keywords,

            cp.comments,

            cp.category,

        ]

    )

    add_result(

        results,

        file_path,

        "Word",

        "C1",

        "プロパティ情報削除",

        "FAIL" if has_props else "PASS",

        "プロパティ情報が残っています。" if has_props else "主要プロパティは空です。",

        "ファイル情報のプロパティを削除。" if has_props else "対応不要。",

    )

 

    if cover_keyword:

        first_page_text = "\n".join(p.text for p in doc.paragraphs[:20])

        cover_ok = cover_keyword in first_page_text

        add_result(

            results,

            file_path,

            "Word",

            "C2",

            "表紙が規定のもの",

            "PASS" if cover_ok else "FAIL",

            f"表紙キーワード '{cover_keyword}' を " + ("検出。" if cover_ok else "未検出。"),

            "規定の表紙へ差し替え。" if not cover_ok else "対応不要。",

        )

    else:

        add_result(

            results,

            file_path,

            "Word",

            "C2",

            "表紙が規定のもの",

            "MANUAL",

            "cover keyword 未指定のため手動確認です。",

            "--cover-keyword を指定すると半自動確認できます。",

        )

 

    margin_fail_msgs = []

    for idx, section in enumerate(doc.sections, start=1):

        top = mm_from_emu(section.top_margin)

        bottom = mm_from_emu(section.bottom_margin)

        left = mm_from_emu(section.left_margin)

        right = mm_from_emu(section.right_margin)

        if top is None or bottom is None or left is None or right is None:

            continue

        if not (top >= 20 and bottom >= 20 and left >= 30 and right >= 20):

            margin_fail_msgs.append(

                f"sec{idx}(上{top:.1f}/下{bottom:.1f}/左{left:.1f}/右{right:.1f}mm)"

            )

 

    add_result(

        results,

        file_path,

        "Word",

        "C3",

        COMMON_MARGIN_RULE_TEXT,

        "FAIL" if margin_fail_msgs else "PASS",

        (

            "; ".join(margin_fail_msgs)

            if margin_fail_msgs

            else "; ".join(

                f"sec{idx}(上{mm_from_emu(section.top_margin):.1f}/下{mm_from_emu(section.bottom_margin):.1f}/左{mm_from_emu(section.left_margin):.1f}/右{mm_from_emu(section.right_margin):.1f}mm)"

                for idx, section in enumerate(doc.sections, start=1)

            )

        ),

        "ページ設定で余白を補正。" if margin_fail_msgs else "対応不要。",

    )

 

    header_footer_xml = "\n".join(

        sec.header._element.xml + sec.footer._element.xml for sec in doc.sections

    )

    has_page_number_field = ("PAGE" in header_footer_xml) or ("w:pgNum" in header_footer_xml)

    add_result(

        results,

        file_path,

        "Word",

        "C4",

        "ページ番号",

        "PASS" if has_page_number_field else "WARN",

        "ページ番号フィールドを検出。" if has_page_number_field else "ページ番号フィールドを検出できません。最大ページ数は画像シート参照。",

        "ヘッダ/フッタにページ番号を設定。" if not has_page_number_field else "対応不要。",

    )

 

    add_result(

        results,

        file_path,

        "Word",

        "C5",

        "PDF出力結果確認（見切れ/罫線/表サイズ/ページ番号）",

        "MANUAL",

        "PDF出力後の品質確認は目視が必要です。",

        "PDFを開き、見切れ・罫線欠け・表サイズ・ページ番号を確認。",

    )

 

    package = getattr(getattr(doc, "part", None), "package", None)

    close_fn = getattr(package, "close", None)

    if callable(close_fn):

        try:

            close_fn()

        except Exception:

            pass

 

def check_pdf(file_path: Path, results: List[CheckResult], cover_keyword: Optional[str]) -> None:

    if PdfReader is None:

        add_result(

            results,

            file_path,

            "PDF",

            "P-ENV",

            "PDF解析ライブラリ",

            "ERROR",

            "pypdf が未インストールのため解析できません。",

            "pip install pypdf を実行してください。"

        )

        return

 

    try:

        reader = PdfReader(str(file_path))

    except Exception as exc:

        add_result(

            results,

            file_path,

            "PDF",

            "P-OPEN",

            "PDFファイル読込",

            "ERROR",

            f"読込失敗: {exc}",

            "ファイル破損・パスワード・権限制限を確認してください。"

        )

        return

 

    metadata = reader.metadata or {}

    meaningful_meta = []

    for key in ["/Author", "/Creator", "/Producer", "/Title", "/Subject", "/Keywords"]:

        val = metadata.get(key)

        if val:

            meaningful_meta.append(f"{key}={val}")

 

    add_result(

        results,

        file_path,

        "PDF",

        "C1",

        "プロパティ情報削除",

        "FAIL" if meaningful_meta else "PASS",

        "; ".join(meaningful_meta) if meaningful_meta else "主要メタデータは空です。",

        "PDF作成時にメタデータ削除して再出力。" if meaningful_meta else "対応不要。"

    )

 

    if cover_keyword and reader.pages:

        first_text = (reader.pages[0].extract_text() or "")[:3000]

        cover_ok = cover_keyword in first_text

        add_result(

            results,

            file_path,

            "PDF",

            "C2",

            "表紙が規定のもの",

            "PASS" if cover_ok else "FAIL",

            f"表紙キーワード '{cover_keyword}' を " + ("検出。" if cover_ok else "未検出。"),

            "規定の表紙へ差し替え。" if not cover_ok else "対応不要。"

        )

    else:

        add_result(

            results,

            file_path,

            "PDF",

            "C2",

            "表紙が規定のもの",

            "MANUAL",

            "cover keyword 未指定または1ページ目取得不可のため手動確認です。",

            "--cover-keyword を指定すると半自動確認できます。",

        )

 

    add_result(

        results,

        file_path,

        "PDF",

        "C3",

        COMMON_MARGIN_RULE_TEXT,

        "MANUAL",

        "PDFはレイアウト固定で余白の厳密自動判定が難しいため手動確認です（基準: 上20/下20/左30/右20mm以上）。",

        "表示倍率100%で余白とパンチ穴干渉を確認。",

    )

 

    page_count = len(reader.pages)

    strikeout_annots = 0

    comment_annots = 0

    other_markup_annots = 0

    strike_pages: List[int] = []

    comment_pages: List[int] = []

    markup_pages: List[int] = []

    ref_error_pages: List[int] = []

    ref_error_pattern = re.compile(r"(?i)(error!\s*reference source not found\.?|#ref!)")

    for page_index, page in enumerate(reader.pages, start=1):

        try:

            page_text = page.extract_text() or ""

        except Exception:

            page_text = ""

        if ref_error_pattern.search(page_text):

            ref_error_pages.append(page_index)

 

        annots = page.get("/Annots")

        if not annots:

            continue

        for annot_ref in annots:

            try:

                annot = annot_ref.get_object()

            except Exception:

                continue

            subtype = annot.get("/Subtype")

            if subtype == "/StrikeOut":

                strikeout_annots += 1

                strike_pages.append(page_index)

            elif subtype in {"/Text", "/FreeText", "/Popup"}:

                comment_annots += 1

                comment_pages.append(page_index)

            elif subtype in {"/Underline", "/Squiggly", "/Highlight", "/Caret"}:

                other_markup_annots += 1

                markup_pages.append(page_index)

 

    add_result(

        results,

        file_path,

        "PDF",

        "P1",

        "見え消し（注釈ベース）残存",

        "FAIL" if (strikeout_annots + other_markup_annots) > 0 else "PASS",

        (

            f"注釈検出: StrikeOut={strikeout_annots}, その他マークアップ={other_markup_annots}。 / 指摘対象ページ：{', '.join(f'P{p}' for p in sorted(set(strike_pages + markup_pages)))}"

            if (strikeout_annots + other_markup_annots) > 0

            else "注釈ベースの見え消しは検出されません。 / 指摘対象ページ：なし"

        ),

        "不要な注釈を削除。テキストに焼き込まれた取り消し線は目視確認。"

        if (strikeout_annots + other_markup_annots) > 0

        else "対応不要。",

    )

 

    add_result(

        results,

        file_path,

        "PDF",

        "P2",

        "コメント（吹き出し）残存",

        "FAIL" if comment_annots > 0 else "PASS",

        (

            f"コメント系注釈: {comment_annots} 件。 / 指摘対象ページ：{', '.join(f'P{p}' for p in sorted(set(comment_pages)))}"

            if comment_annots > 0

            else "コメント系注釈: 0 件。 / 指摘対象ページ：なし"

        ),

        "不要なコメント注釈を削除。" if comment_annots > 0 else "対応不要。",

    )

 

    add_result(

        results,

        file_path,

        "PDF",

        "P3",

        "参照エラー残存",

        "FAIL" if ref_error_pages else "PASS",

        (

            f"参照エラー文字列を検出。 / 指摘対象ページ：{', '.join(f'P{p}' for p in sorted(set(ref_error_pages)))}"

            if ref_error_pages

            else "参照エラー文字列は検出されません。 / 指摘対象ページ：なし"

        ),

        "参照元を修正し、PDF再出力。" if ref_error_pages else "対応不要。",

    )

 

    add_result(

        results,

        file_path,

        "PDF",

        "C4",

        "ページ番号",

        "MANUAL",

        f"ページ数: {page_count}。抽出品質の都合でページ番号整合は手動確認を推奨。",

        "全ページのページ番号連番を確認。",

    )

 

    add_result(

        results,

        file_path,

        "PDF",

        "C5",

        "PDF出力結果確認（見切れ/罫線/表サイズ/ページ番号）",

        "MANUAL",

        "PDF品質確認の総合判定は目視が必要です。詳細は C5.1?C5.4 を参照。",

        "C5.1?C5.4 を確認して総合判断。",

    )

 

    add_result(

        results,

        file_path,

        "PDF",

        "C5.1",

        "PDF見切れ",

        "MANUAL",

        "見切れは表示環境依存のため目視確認が必要です。",

        "各ページ端部の文字欠けを確認。",

    )

    add_result(

        results,

        file_path,

        "PDF",

        "C5.2",

        "PDF罫線欠け",

        "MANUAL",

        "罫線欠けの判定は目視確認が必要です。",

        "細線・表罫線の欠けを確認。",

    )

    add_result(

        results,

        file_path,

        "PDF",

        "C5.3",

        "PDF表サイズ極小",

        "MANUAL",

        "表サイズの妥当性は文書意図依存のため目視確認です。",

        "表が過度に縮小されていないか確認。",

    )

    add_result(

        results,

        file_path,

        "PDF",

        "C5.4",

        "PDFページ番号正当性",

        "MANUAL",

        "ページ番号表示と総ページ整合は目視確認です。",

        "ページ番号の欠番・重複を確認。",

    )

 

    reader_stream = getattr(reader, "stream", None)

    if reader_stream is not None:

        try:

            reader_stream.close()

        except Exception:

            pass

 

def check_file(file_path: Path, results: List[CheckResult], cover_keyword: Optional[str]) -> None:

    suffix = file_path.suffix.lower()

    file_type = ""

    if suffix in {".xlsx", ".xlsm"}:

        file_type = "Excel"

        check_excel(file_path, results)

    elif suffix == ".docx":

        file_type = "Word"

        check_word(file_path, results, cover_keyword)

    elif suffix == ".pdf":

        file_type = "PDF"

        check_pdf(file_path, results, cover_keyword)

    elif suffix in {".xls", ".doc"}:

        file_type = "LegacyOffice"

        add_result(

            results,

            file_path,

            "LegacyOffice",

            "L1",

            "旧形式ファイル",

            "WARN",

            "旧形式(.xls/.doc)は本スクリプトでの詳細解析対象外です。",

            "可能なら .xlsx/.docx へ変換して再チェック。",

        )

 

    if file_type:

        ensure_expected_checks(results, file_path, file_type)

 

def find_target_files(root: Path, exclude_paths: Optional[Set[Path]] = None) -> List[Path]:

    target_suffixes = {".xlsx", ".xlsm", ".xls", ".docx", ".doc", ".pdf"}

    found: List[Path] = []

    exclude_resolved: Set[Path] = {p.resolve() for p in (exclude_paths or set())}

    ignored_dir_names = {".venv", "venv", ".git", "__pycache__"}

    for dirpath, dirnames, filenames in os.walk(root):

        dirnames[:] = [d for d in dirnames if d not in ignored_dir_names]

        dir_path = Path(dirpath)

        resolved_dir = dir_path.resolve()

        if resolved_dir in exclude_resolved or any(parent in exclude_resolved for parent in resolved_dir.parents):

            dirnames[:] = []

            continue

        for filename in filenames:

            if filename.startswith("~$"):

                continue

            path = dir_path / filename

            if path.suffix.lower() not in target_suffixes:

                continue

            resolved = path.resolve()

            if resolved in exclude_resolved:

                continue

            if any(parent in exclude_resolved for parent in resolved.parents):

                continue

            found.append(path)

    return sorted(set(found))

 

def main() -> None:

    start_total = time.perf_counter()

    parser = argparse.ArgumentParser(

        description="指定フォルダ配下のExcel/Word/PDFをレビュー基準でチェックし、xlsx一覧を出力します。"

    )

    parser.add_argument("root_folder", help="チェック対象のルートフォルダ")

    parser.add_argument(

        "-o",

        "--output",

        default="review_results.xlsx",

        help="出力xlsxパス (既定: review_results.xlsx)",

    )

    parser.add_argument(

        "--cover-keyword",

        default=None,

        help="表紙判定用キーワード（任意）。指定時のみ半自動で表紙チェック。",

    )

    parser.add_argument(

        "--visual-assets-dir",

        default="review_visual_assets",

        help="PDF/PNG生成物の保存先フォルダ (既定: review_visual_assets)",

    )

    parser.add_argument(

        "--preview-pages-per-row",

        type=int,

        default=6,

        help="互換用オプション（現在はimage_previewを横スクロール固定で全ページ表示）",

    )

    parser.add_argument(

        "--no-visual",

        action="store_true",

        help="画像生成（PDF/PNG・image_preview）を無効化する",

    )

    args = parser.parse_args()

 

    root = Path(args.root_folder).resolve()

    out_xlsx = Path(args.output).resolve()

 

    if not root.exists() or not root.is_dir():

        raise SystemExit(f"対象フォルダが存在しません: {root}")

 

    assets_root = Path(args.visual_assets_dir).resolve()

    target_files = find_target_files(root, exclude_paths={out_xlsx, assets_root})

    results: List[CheckResult] = []

    visual_pages: List[VisualPage] = []

    visual_enabled = not args.no_visual

 

    if not target_files:

        add_result(

            results,

            root,

            "System",

            "S1",

            "対象ファイル検出",

            "WARN",

            "対象ファイル（Excel/Word/PDF）が見つかりませんでした。",

            "フォルダ構成と拡張子を確認。",

        )

    else:

        for idx, file_path in enumerate(target_files, start=1):

            start_file = time.perf_counter()

            print(f"[{idx}/{len(target_files)}] チェック開始: {file_path}")

            try:

                check_file(file_path, results, args.cover_keyword)

                page_count = None

                if visual_enabled:

                    page_count = run_visual_pipeline(file_path, results, visual_pages, assets_root)

                append_max_page_detail(results, file_path, page_count)

            except Exception as exc:

                add_result(

                    results,

                    file_path,

                    "System",

                    "S3",

                    "ファイル処理例外",

                    "ERROR",

                    f"処理中に例外が発生しました: {exc}",

                    "対象ファイルを個別確認し、必要なら単体再実行してください。",

                )

            finally:

                if (idx % 20) == 0:

                    gc.collect()

            elapsed_file = time.perf_counter() - start_file

            print(f"[{idx}/{len(target_files)}] チェック完了: {file_path} ({elapsed_file:.2f} 秒)")

 

    gc.collect()

 

    err = write_visual_report_xlsx(

        results,

        visual_pages,

        out_xlsx,

        preview_pages_per_row=args.preview_pages_per_row,

    )

    if err:

        results_path = out_xlsx.with_name(f"{out_xlsx.stem}_results{out_xlsx.suffix}")

        images_path = out_xlsx.with_name(f"{out_xlsx.stem}_image_preview{out_xlsx.suffix}")

        results_err = write_results_report_xlsx(results, results_path)

        images_err = write_image_preview_xlsx(

            visual_pages,

            images_path,

            preview_pages_per_row=args.preview_pages_per_row,

        )

        if results_err or images_err:

            raise SystemExit(

                "xlsx出力失敗: "

                + err

                + " / 分割出力失敗: "

                + ", ".join(e for e in [results_err, images_err] if e)

            )

        print(f"通常xlsx出力失敗のため分割出力にフォールバック: {results_path}, {images_path}")

 

    total = len(results)

    fail = sum(1 for r in results if r.status == "FAIL")

    warn = sum(1 for r in results if r.status == "WARN")

    manual = sum(1 for r in results if r.status == "MANUAL")

    error = sum(1 for r in results if r.status == "ERROR")

    na = sum(1 for r in results if r.status == "N/A")

 

    print(f"対象ファイル数: {len(target_files)}")

    print(

        f"チェック結果数: {total} "

        f"(FAIL={fail}, WARN={warn}, MANUAL={manual}, ERROR={error}, N/A={na})"

    )

    print(f"xlsx出力: {out_xlsx}")

    print(f"画像生成: {'有効' if visual_enabled else '無効'}")

    elapsed_total = time.perf_counter() - start_total

    print(f"処理時間: {elapsed_total:.2f} 秒")

 

if __name__ == "__main__":

    main()
